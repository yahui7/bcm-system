"""
BCP 生成 + Word 导出引擎
"""
import json
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class BCPEngine:
    def __init__(self, business_id: int):
        self.business_id = business_id

    def generate(self, conn) -> dict:
        """自动生成 BCP JSON"""
        cursor = conn.cursor()

        # 业务信息
        cursor.execute("SELECT * FROM businesses WHERE id = ?", (self.business_id,))
        biz = cursor.fetchone()
        if not biz:
            return {}

        # 资源列表
        cursor.execute("SELECT * FROM resources WHERE business_id = ?", (self.business_id,))
        resources = cursor.fetchall()

        # 风险场景
        cursor.execute(
            "SELECT * FROM risk_scenarios WHERE business_id = ? AND risk_level IN ('high', 'medium') ORDER BY risk_score DESC LIMIT 10",
            (self.business_id,),
        )
        scenarios = cursor.fetchall()

        # BIA 信息
        cursor.execute("SELECT * FROM bia_factor_scores WHERE business_id = ?", (self.business_id,))
        bia_scores = cursor.fetchall()

        # 策略库
        cursor.execute("SELECT * FROM recovery_strategies WHERE enabled = 1")
        strategies = cursor.fetchall()

        # 组装 BCP
        bcp = {
            "bcp_name": f"{biz['name']}业务连续性计划",
            "version": "1.0",
            "plan_info": {
                "owner": biz["owner"],
                "department": biz["department"],
                "contact": "",
                "last_updated": biz["updated_at"] or "",
            },
            "critical_processes": [
                {"name": biz["name"], "mad": "< 4小时" if biz["bia_tier"] == "Tier 1" else "< 24小时",
                 "rto": "< 2小时" if biz["bia_tier"] == "Tier 1" else "< 8小时",
                 "rpo": "Near-Zero" if biz["bia_tier"] == "Tier 1" else "< 24小时"},
            ],
            "preventive_strategies": [s["name"] for s in strategies if s["strategy_type"] == "preventive"][:3],
            "recovery_strategies": [s["name"] for s in strategies if s["strategy_type"] == "recovery"][:3],
            "rep": [],
            "war_location": "",
            "contacts": [],
            "recovery_actions": [
                {"scenario": s["threat"], "priority": i + 1, "action": "启动对应恢复策略", "responsible": biz["owner"]}
                for i, s in enumerate(scenarios[:4])
            ],
        }

        return bcp


def export_bcp_word(bcp_data: dict) -> BytesIO:
    """生成 BCP Word 文档"""
    doc = Document()

    # 标题页
    title = doc.add_heading("业务连续性计划（BCP）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"业务名称：{bcp_data.get('bcp_name', '')}").bold = True

    doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"版本：{bcp_data.get('version', '1.0')}")

    plan_info = bcp_data.get("plan_info", {})
    doc.add_paragraph("")

    # 1. 计划信息
    doc.add_heading("1. 计划信息", level=1)
    table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    info_rows = [
        ("负责人", plan_info.get("owner", "")),
        ("部门", plan_info.get("department", "")),
        ("联系方式", plan_info.get("contact", "")),
        ("更新日期", plan_info.get("last_updated", "")),
    ]
    for i, (k, v) in enumerate(info_rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    # 2. 关键业务流程
    doc.add_heading("2. 关键业务流程", level=1)
    proc_table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
    hdr = proc_table.rows[0].cells
    hdr[0].text = "流程名称"; hdr[1].text = "MAD"; hdr[2].text = "RTO"; hdr[3].text = "RPO"
    for proc in bcp_data.get("critical_processes", []):
        row = proc_table.add_row().cells
        row[0].text = proc.get("name", ""); row[1].text = proc.get("mad", "")
        row[2].text = proc.get("rto", ""); row[3].text = proc.get("rpo", "")

    # 3. 恢复策略
    doc.add_heading("3. 恢复策略", level=1)
    doc.add_heading("预防策略", level=2)
    for s in bcp_data.get("preventive_strategies", []):
        doc.add_paragraph(s, style="List Bullet")
    doc.add_heading("恢复策略", level=2)
    for s in bcp_data.get("recovery_strategies", []):
        doc.add_paragraph(s, style="List Bullet")

    # 4. REP
    doc.add_heading("4. 恢复关键人员（REP）", level=1)
    rep_table = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
    hdr = rep_table.rows[0].cells
    hdr[0].text = "姓名"; hdr[1].text = "角色"; hdr[2].text = "电话"
    for rep in bcp_data.get("rep", []):
        row = rep_table.add_row().cells
        row[0].text = rep.get("name", ""); row[1].text = rep.get("role", "")
        row[2].text = rep.get("phone", "")

    # 5. WAR
    doc.add_heading("5. 备用工作场所（WAR）", level=1)
    doc.add_paragraph(bcp_data.get("war_location", "待填写"))

    # 6. 关键联系人
    doc.add_heading("6. 关键联系人", level=1)
    contact_table = doc.add_table(rows=1, cols=3, style="Light Grid Accent 1")
    hdr = contact_table.rows[0].cells
    hdr[0].text = "姓名/单位"; hdr[1].text = "类型"; hdr[2].text = "电话"
    for c in bcp_data.get("contacts", []):
        row = contact_table.add_row().cells
        row[0].text = c.get("name", ""); row[1].text = c.get("type", "")
        row[2].text = c.get("phone", "")

    # 7. 恢复行动步骤
    doc.add_heading("7. 恢复行动步骤", level=1)
    act_table = doc.add_table(rows=1, cols=4, style="Light Grid Accent 1")
    hdr = act_table.rows[0].cells
    hdr[0].text = "场景"; hdr[1].text = "优先级"; hdr[2].text = "行动"; hdr[3].text = "负责人"
    for act in bcp_data.get("recovery_actions", []):
        row = act_table.add_row().cells
        row[0].text = act.get("scenario", "")
        row[1].text = str(act.get("priority", ""))
        row[2].text = act.get("action", "")
        row[3].text = act.get("responsible", "")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
