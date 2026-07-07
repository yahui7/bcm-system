"""
演练管理 API
"""
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from io import BytesIO
from docx import Document
from backend.database import get_connection

router = APIRouter()


@router.get("/plans")
async def list_plans():
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """SELECT d.*, b.name as business_name
           FROM drill_plans d
           LEFT JOIN businesses b ON d.business_id = b.id
           ORDER BY d.drill_date DESC"""
    )
    rows = cursor.fetchall()
    conn.close()
    return {"status": "ok", "plans": [dict(r) for r in rows]}


class PlanCreate(BaseModel):
    business_id: int
    drill_date: str = ""
    drill_type: str = "桌面推演"
    participants: str = ""
    objective: str = ""


@router.post("/plans")
async def create_plan(data: PlanCreate):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO drill_plans (business_id, drill_date, drill_type, participants, objective) VALUES (?, ?, ?, ?, ?)",
        (data.business_id, data.drill_date, data.drill_type, data.participants, data.objective),
    )
    conn.commit()
    pid = cursor.lastrowid
    conn.close()
    return {"status": "ok", "id": pid, "message": "演练计划已创建"}


@router.put("/plans/{plan_id}")
async def update_plan(plan_id: int, data: dict):
    conn = get_connection()
    cursor = conn.cursor()
    allowed = ["drill_date", "drill_type", "participants", "objective", "business_id", "status"]
    updates = {k: data[k] for k in data if k in allowed}
    if not updates:
        conn.close()
        raise HTTPException(status_code=400, detail="无有效更新字段")
    set_clause = ", ".join(f"{k} = ?" for k in updates)
    values = list(updates.values()) + [plan_id]
    cursor.execute(f"UPDATE drill_plans SET {set_clause} WHERE id = ?", values)
    conn.commit()
    conn.close()
    return {"status": "ok", "message": "更新成功"}


@router.delete("/plans/{plan_id}")
async def delete_plan(plan_id: int):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("DELETE FROM drill_plans WHERE id = ?", (plan_id,))
    conn.commit()
    conn.close()
    return {"status": "ok", "message": "已删除"}


@router.get("/results")
async def list_results():
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM drill_results ORDER BY completed_at DESC")
    rows = cursor.fetchall()
    conn.close()
    return {"status": "ok", "results": [dict(r) for r in rows]}


class ResultSubmit(BaseModel):
    passed: int = 1
    failure_reason: str = ""
    actual_recovery_time: str = ""


@router.post("/plans/{plan_id}/result")
async def submit_result(plan_id: int, data: ResultSubmit):
    conn = get_connection()
    cursor = conn.cursor()

    # 检查是否已有结果
    cursor.execute("SELECT id FROM drill_results WHERE plan_id = ?", (plan_id,))
    existing = cursor.fetchone()
    if existing:
        raise HTTPException(status_code=400, detail="该计划已有演练结果")

    cursor.execute(
        "INSERT INTO drill_results (plan_id, passed, failure_reason, actual_recovery_time) VALUES (?, ?, ?, ?)",
        (plan_id, data.passed, data.failure_reason, data.actual_recovery_time),
    )

    # 更新计划状态
    cursor.execute("UPDATE drill_plans SET status = '已完成' WHERE id = ?", (plan_id,))

    conn.commit()
    conn.close()
    return {"status": "ok", "message": "演练结果已记录"}


@router.get("/plans/{plan_id}/export")
async def export_result(plan_id: int):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM drill_plans WHERE id = ?", (plan_id,))
    plan = cursor.fetchone()
    cursor.execute("SELECT * FROM drill_results WHERE plan_id = ?", (plan_id,))
    result = cursor.fetchone()
    conn.close()

    if not plan:
        raise HTTPException(status_code=404, detail="计划不存在")

    doc = Document()
    doc.add_heading("演练报告", level=0)
    doc.add_paragraph(f"演练计划 ID：{plan['id']}")
    doc.add_paragraph(f"演练日期：{plan['drill_date']}")
    doc.add_paragraph(f"演练类型：{plan['drill_type']}")
    doc.add_paragraph(f"参与人员：{plan['participants']}")
    doc.add_paragraph(f"预期目标：{plan['objective']}")

    if result:
        doc.add_heading("演练结果", level=1)
        doc.add_paragraph(f"结果：{'✅ 通过' if result['passed'] else '❌ 失败'}")
        if result["failure_reason"]:
            doc.add_paragraph(f"失败原因：{result['failure_reason']}")
        doc.add_paragraph(f"实际恢复时间：{result['actual_recovery_time']}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=drill_report_{plan_id}.docx"},
    )
