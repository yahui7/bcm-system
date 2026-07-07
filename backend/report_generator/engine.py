"""
报告生成引擎 — BCM 专用
"""
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.database import get_connection


class ReportEngine:
    """报告生成引擎"""

    def __init__(self):
        pass

    # ── 数据源 ──

    def _get_bia_summary(self) -> dict:
        """BIA 业务影响分析数据，返回占位符字典"""
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM businesses ORDER BY id")
        businesses = cursor.fetchall()

        if not businesses:
            conn.close()
            return {"业务概况": "暂无业务数据", "评分矩阵": "暂无", "等级结果": "暂无", "依赖分析": "暂无", "结论": "暂无法评估"}

        # 业务概况
        biz_lines = []
        for b in businesses:
            biz_lines.append(f"· {b['name']}（{b['department']}）— {b['description'][:60] if b['description'] else '无'}")

        # 评分矩阵
        matrix_lines = []
        for b in businesses:
            cursor.execute(
                """SELECT s.*, f.name as factor_name
                   FROM bia_factor_scores s
                   JOIN bia_factors f ON s.factor_id = f.id
                   WHERE s.business_id = ?""",
                (b["id"],),
            )
            scores = cursor.fetchall()
            if scores:
                matrix_lines.append(f"\n{b['name']}：")
                for s in scores:
                    matrix_lines.append(f"  {s['factor_name']}：D1={s['day1']} D3={s['day3']} W1={s['week1']} M1={s['month1']}")

        # 等级结果
        tier_lines = []
        for b in businesses:
            tier = b["bia_tier"] or "未评级"
            level = {"Tier 1": "关键业务", "Tier 2": "重要业务", "Tier 3": "一般业务"}.get(tier, tier)
            tier_lines.append(f"· {b['name']}：{tier}（{level}） 综合评分 {b['bia_score'] or 0} 分")

        # 依赖分析
        dep_lines = []
        for b in businesses:
            cursor.execute(
                """SELECT d.*, b2.name as related_name
                   FROM business_dependencies d
                   JOIN businesses b2 ON d.related_business_id = b2.id
                   WHERE d.business_id = ?""",
                (b["id"],),
            )
            deps = cursor.fetchall()
            if deps:
                dep_lines.append(f"\n{b['name']} 的依赖关系：")
                for d in deps:
                    t = "上游（我依赖它）" if d["dependency_type"] == "upstream" else "下游（它依赖我）"
                    dep_lines.append(f"  → {d['related_name']}（{t}）")

        # 结论
        tier1_count = sum(1 for b in businesses if b["bia_tier"] == "Tier 1")
        conclusion = f"共评估 {len(businesses)} 个业务，其中关键业务（Tier 1）{tier1_count} 个"
        if tier1_count > 0:
            conclusion += "，需重点关注并制定完整的 BCP 计划"

        conn.close()
        return {
            "业务概况": "\n".join(biz_lines),
            "评分矩阵": "\n".join(matrix_lines) if matrix_lines else "暂无评分数据",
            "等级结果": "\n".join(tier_lines),
            "依赖分析": "\n".join(dep_lines) if dep_lines else "暂无上下游依赖关系",
            "结论": conclusion,
        }

    def _get_risk_assessment(self) -> dict:
        """RA 风险评估数据"""
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM businesses WHERE is_critical = 1 OR bia_tier IN ('Tier 1', 'Tier 2')")
        businesses = cursor.fetchall()

        if not businesses:
            conn.close()
            return {"资源统计": "暂无重要业务数据", "资源清单": "暂无", "场景总数": "0", "场景分析": "暂无", "高风险汇总": "暂无", "整改建议": "暂无高风险项"}

        # 资源统计和清单
        total_resources = 0
        res_lines = []
        for b in businesses:
            cursor.execute("SELECT * FROM resources WHERE business_id = ?", (b["id"],))
            resources = cursor.fetchall()
            total_resources += len(resources)
            if resources:
                res_lines.append(f"\n{b['name']}（{b['bia_tier'] or '未评级'}）：")
                for r in resources:
                    imp = {"high": "高", "medium": "中", "low": "低"}.get(r["importance"], r["importance"])
                    res_lines.append(f"  [{r['resource_type']}] {r['name']}  重要程度：{imp}  RTO：{r['rto']}  RPO：{r['rpo']}")

        resource_stat = f"共评估 {len(businesses)} 个重要业务，识别 {total_resources} 个关键资源"

        # 风险场景分析
        total_scenarios = 0
        high_scenarios = []
        scenario_lines = []
        for b in businesses:
            cursor.execute(
                "SELECT * FROM risk_scenarios WHERE business_id = ? ORDER BY risk_score DESC",
                (b["id"],),
            )
            scenarios = cursor.fetchall()
            total_scenarios += len(scenarios)
            biz_high = [s for s in scenarios if s["risk_level"] == "high"]
            high_scenarios.extend(biz_high)
            biz_mid = [s for s in scenarios if s["risk_level"] == "medium"]
            scenario_lines.append(
                f"\n{b['name']}：共 {len(scenarios)} 个场景（高风险 {len(biz_high)}、中风险 {len(biz_mid)}、"
                f"低风险 {len(scenarios) - len(biz_high) - len(biz_mid)}）"
            )

        # 高风险汇总
        high_lines = []
        if high_scenarios:
            high_scenarios.sort(key=lambda s: s["risk_score"], reverse=True)
            high_lines.append(f"共发现 {len(high_scenarios)} 个高风险场景，TOP 5：")
            for s in high_scenarios[:5]:
                high_lines.append(
                    f"  · {s['resource_name']} × {s['threat']}：风险值 {s['risk_score']:.0f}"
                    f"（威胁 {s['threat_score']} × 脆弱性 {s['vulnerability_score']}）"
                )

        # 整改建议
        if high_scenarios:
            rec_lines = ["针对高风险场景，建议："]
            rec_lines.append(f"1. 优先处理 {len(high_scenarios)} 个高风险场景，制定风险缓解措施")
            rec_lines.append("2. 对高风险资源增加冗余备份或增强防护")
            rec_lines.append("3. 将高风险项纳入 BCP 计划，制定具体恢复策略")
            rec_lines.append("4. 定期重新评估，跟踪风险变化趋势")
        else:
            rec_lines = ["未发现高风险场景，当前风险可控。建议保持持续监控。"]

        conn.close()
        return {
            "资源统计": resource_stat,
            "资源清单": "\n".join(res_lines),
            "场景总数": str(total_scenarios),
            "场景分析": "\n".join(scenario_lines),
            "高风险汇总": "\n".join(high_lines) if high_lines else "未发现高风险场景",
            "整改建议": "\n".join(rec_lines),
        }

    def _get_bcp_summary(self) -> dict:
        """BCP 连续性计划摘要"""
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            """SELECT p.*, b.name as business_name
               FROM bcp_plans p
               JOIN businesses b ON p.business_id = b.id
               ORDER BY p.id"""
        )
        plans = cursor.fetchall()
        conn.close()

        if not plans:
            return {"BCP概况": "暂无 BCP 数据"}
        lines = [f"共制定 {len(plans)} 份业务连续性计划："]
        for p in plans:
            lines.append(f"· {p['business_name']} — {p['plan_name']}（{p['status']}）")
        return {"BCP概况": "\n".join(lines)}

    def _get_drill_summary(self) -> dict:
        """演练结果数据"""
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            """SELECT d.*, b.name as business_name
               FROM drill_plans d
               LEFT JOIN businesses b ON d.business_id = b.id
               ORDER BY d.drill_date DESC"""
        )
        plans = cursor.fetchall()

        if not plans:
            conn.close()
            return {"演练基本信息": "暂无演练数据", "演练过程": "暂无", "演练结果": "暂无", "改进建议": "暂无"}

        # 基本信息
        info_lines = [f"共制定 {len(plans)} 个演练计划"]

        # 演练过程
        process_lines = []
        for p in plans:
            process_lines.append(f"\n· {p['business_name']} — {p['drill_type']}")
            process_lines.append(f"  日期：{p['drill_date']}  参与人员：{p['participants']}")
            process_lines.append(f"  目标：{p['objective']}")

        # 演练结果
        result_lines = []
        passed = 0
        for p in plans:
            cursor.execute("SELECT * FROM drill_results WHERE plan_id = ?", (p["id"],))
            r = cursor.fetchone()
            if r:
                if r["passed"]:
                    passed += 1
                    result_lines.append(f"· {p['business_name']}：✅ 通过（恢复时间 {r['actual_recovery_time']}）")
                else:
                    result_lines.append(f"· {p['business_name']}：❌ 失败（{r['failure_reason']}）")
            else:
                result_lines.append(f"· {p['business_name']}：未记录结果")

        # 改进建议
        completed = sum(1 for p in plans if p["status"] == "已完成")
        improve_lines = []
        if completed < len(plans):
            improve_lines.append(f"1. 尚有 {len(plans) - completed} 个演练未完成，建议尽快安排")
        if passed < completed:
            improve_lines.append(f"2. {completed - passed} 个演练未通过，需分析原因并重新演练")
        improve_lines.append("3. 建议每季度至少开展一次桌面推演，每年至少一次实战演练")
        improve_lines.append("4. 演练结果应纳入 BCP 持续改进流程")

        conn.close()
        return {
            "演练基本信息": "\n".join(info_lines),
            "演练过程": "\n".join(process_lines),
            "演练结果": "\n".join(result_lines),
            "改进建议": "\n".join(improve_lines),
        }

    # ── 生成报告 ──

    def generate(self, template: dict, info: dict) -> dict:
        """根据模板生成报告内容"""
        source_handlers = {
            "bia_summary": self._get_bia_summary,
            "risk_assessment": self._get_risk_assessment,
            "bcp_summary": self._get_bcp_summary,
            "drill_summary": self._get_drill_summary,
        }

        # 预设模板 → 占位符替换
        if template.get("preset"):
            # 确定数据源类型
            tpl_id = template["id"]
            preset_sources = {
                "tpl_bia": "bia_summary",
                "tpl_risk": "risk_assessment",
                "tpl_drill": "drill_summary",
            }
            source_key = preset_sources.get(tpl_id, "bia_summary")
            handler = source_handlers.get(source_key)

            # 获取占位符数据
            placeholders = {}
            if handler:
                placeholders = handler()

            # 默认占位符
            placeholders.setdefault("公司名称", "XX银行")
            placeholders.setdefault("评估日期", info.get("date", datetime.now().strftime("%Y-%m-%d")))
            placeholders.setdefault("编制部门", info.get("author", "风险管理部"))

            # 替换
            sections = []
            for sec in template.get("sections", []):
                content = sec.get("content", "")
                for key, val in placeholders.items():
                    content = content.replace("{{" + key + "}}", val)
                sections.append({"title": sec["title"], "source": source_key, "content": content})

        else:
            # 自定义模板 → 调用 handler
            sections = []
            for sec in template.get("sections", []):
                source = sec.get("source", "manual")
                content = ""
                if source in source_handlers:
                    data = source_handlers[source]()
                    content = "\n".join(data.values()) if isinstance(data, dict) else data
                sections.append({"title": sec["title"], "source": source, "content": content})

        return {
            "title": info.get("title", ""),
            "report_no": info.get("report_no", ""),
            "author": info.get("author", ""),
            "date": info.get("date", datetime.now().strftime("%Y-%m-%d")),
            "sections": sections,
        }

    # ── 导出 Word ──

    def export_word(self, report: dict) -> BytesIO:
        """将报告导出为 Word 文档"""
        doc = Document()

        title = doc.add_heading(report.get("title", "报告"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_text = f"报告编号：{report.get('report_no', '')}　　编制人：{report.get('author', '')}　　日期：{report.get('date', '')}"
        info_p = doc.add_paragraph(info_text)
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in info_p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        for sec in report.get("sections", []):
            doc.add_heading(sec["title"], level=1)
            content = sec.get("content", "")
            if content:
                for line in content.split("\n"):
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(11)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
